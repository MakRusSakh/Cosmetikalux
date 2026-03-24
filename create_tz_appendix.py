#!/usr/bin/env python3
"""
Генерация Приложения №1 (Техническое задание) к договору разработки сайта Cosmetikalux.
Пакет «Базовый интернет-магазин» — 250 000 ₽
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ══════════════════════════════════════
# СТИЛИ ДОКУМЕНТА
# ══════════════════════════════════════

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)
style.paragraph_format.line_spacing = 1.15

# Поля страницы
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)


def add_heading_centered(text, level=1, bold=True, size=None):
    """Добавить заголовок по центру."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = size or (Pt(16) if level == 0 else Pt(14) if level == 1 else Pt(13))
    p.paragraph_format.space_before = Pt(12) if level <= 1 else Pt(8)
    p.paragraph_format.space_after = Pt(8) if level <= 1 else Pt(4)
    return p


def add_heading_left(text, level=2, bold=True, size=None):
    """Добавить заголовок слева."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = size or (Pt(13) if level == 2 else Pt(12))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_cm=None):
    """Добавить параграф."""
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if indent_cm:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    return p


def add_bullet(text, level=0):
    """Добавить пункт маркированного списка."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p


def set_cell_shading(cell, color):
    """Установить заливку ячейки."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def make_table(headers, rows, col_widths=None):
    """Создать отформатированную таблицу."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Заголовки
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')

    # Данные
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

    # Ширина столбцов
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


# ══════════════════════════════════════
# ТИТУЛЬНАЯ ЧАСТЬ
# ══════════════════════════════════════

add_heading_centered('ПРИЛОЖЕНИЕ №1', level=0, size=Pt(16))
add_heading_centered('к Договору на оказание услуг по разработке интернет-сайта', level=1, size=Pt(13), bold=False)
add_heading_centered(f'от «___» ____________ 2026 г.', level=1, size=Pt(13), bold=False)

doc.add_paragraph()  # пустая строка

add_heading_centered('ТЕХНИЧЕСКОЕ ЗАДАНИЕ', level=0, size=Pt(16))
add_heading_centered('на разработку интернет-магазина cosmetikalux.ru', level=1, size=Pt(14))

doc.add_paragraph()

add_para('Заказчик: ИП Тимошенко Анна Вячеславовна (ИНН 650116762461)', bold=True)
add_para('Исполнитель: ИП Макуха Руслан Александрович (ИНН 650120178498)', bold=True)

doc.add_paragraph()

# ══════════════════════════════════════
# 1. ОБЩИЕ СВЕДЕНИЯ
# ══════════════════════════════════════

add_heading_centered('1. ОБЩИЕ СВЕДЕНИЯ', level=1)

add_heading_left('1.1. Наименование проекта')
add_para('Интернет-магазин корейской и японской косметики «CosmetikaLux» (cosmetikalux.ru).', indent_cm=1.25)

add_heading_left('1.2. Цель проекта')
add_para(
    'Разработка полнофункционального интернет-магазина, заменяющего текущий сайт '
    'на платформе Nethouse, обеспечивающего удобную навигацию по каталогу, '
    'оформление и оплату заказов, управление контентом через административную панель.',
    indent_cm=1.25
)

add_heading_left('1.3. Целевая аудитория')
add_bullet('Основная: женщины 25–45 лет, интересующиеся K-beauty, средний+ доход')
add_bullet('Вторичная: девушки 18–25 лет, начинающие уход за кожей')
add_bullet('Третичная: женщины 45+, ищущие anti-age решения премиум-класса')

add_heading_left('1.4. География')
add_bullet('Основной регион: Южно-Сахалинск и Сахалинская область')
add_bullet('Потенциально: вся Россия (при подключении модуля доставки)')

add_heading_left('1.5. Параметры каталога')
make_table(
    ['Параметр', 'Значение'],
    [
        ['Текущее количество товаров (SKU)', '~478'],
        ['Планируемый рост', 'до 500–1000 SKU'],
        ['Категории', '15–25'],
        ['Бренды', '20–40'],
    ],
    col_widths=[10, 7]
)

# ══════════════════════════════════════
# 2. ПАКЕТ УСЛУГ
# ══════════════════════════════════════

add_heading_centered('2. СОСТАВ РАБОТ (ПАКЕТ «БАЗОВЫЙ ИНТЕРНЕТ-МАГАЗИН»)', level=1)

add_para('Стоимость пакета: 250 000 (двести пятьдесят тысяч) рублей.', bold=True, indent_cm=1.25)

doc.add_paragraph()

add_heading_left('2.1. Каталог и товары')
add_bullet('Каталог товаров с фильтрацией (категория, бренд, цена, тип кожи)')
add_bullet('Страница товара: галерея фото, описание, состав (INCI), объём, способ применения')
add_bullet('Страницы категорий с описанием и товарами')
add_bullet('Страницы брендов с логотипом и описанием')
add_bullet('Полнотекстовый поиск по товарам')
add_bullet('Сортировка (по цене, новизне, популярности)')
add_bullet('Бейджи: «Хит», «Новинка», «Скидка»')

add_heading_left('2.2. Корзина и оформление заказа')
add_bullet('Корзина: добавление, изменение количества, удаление товаров')
add_bullet('Мини-корзина в шапке сайта')
add_bullet('Оформление заказа: контактные данные, адрес доставки')
add_bullet('Онлайн-оплата через ЮKassa (банковская карта, СБП)')
add_bullet('Страница успешной оплаты с номером заказа')

add_heading_left('2.3. Личный кабинет покупателя')
add_bullet('Регистрация и вход (телефон + пароль)')
add_bullet('Профиль: имя, телефон, email')
add_bullet('История заказов с детализацией и статусами')
add_bullet('Избранные товары')
add_bullet('Адреса доставки (сохранение нескольких адресов)')

add_heading_left('2.4. Административная панель')
add_bullet('CRUD товаров (добавление, редактирование, удаление, массовый импорт)')
add_bullet('Управление категориями и брендами')
add_bullet('Управление заказами: просмотр, смена статуса, фильтрация')
add_bullet('Дашборд: заказы за сегодня, выручка, количество товаров')
add_bullet('Управление клиентами: просмотр, поиск')

add_heading_left('2.5. SEO-оптимизация')
add_bullet('Мета-теги (title, description) на всех страницах')
add_bullet('Schema.org: Product, Offer, BreadcrumbList, Organization, WebSite')
add_bullet('Динамический sitemap.xml')
add_bullet('robots.txt')
add_bullet('Хлебные крошки на всех страницах')
add_bullet('Семантическая HTML-вёрстка')
add_bullet('Оптимизация изображений (WebP, lazy loading)')

add_heading_left('2.6. Дизайн и вёрстка')
add_bullet('Уникальный премиальный дизайн в стиле CosmetikaLux')
add_bullet('Адаптивная вёрстка: мобильный, планшет, десктоп')
add_bullet('Главная страница: hero-баннер, хиты продаж, новинки, бренды, преимущества')
add_bullet('Анимации и микроинтеракции')

add_heading_left('2.7. Информационные страницы')
add_bullet('О магазине')
add_bullet('Доставка и оплата')
add_bullet('Возврат и обмен')
add_bullet('Контакты (с картой)')
add_bullet('Политика конфиденциальности')
add_bullet('Пользовательское соглашение')

# ══════════════════════════════════════
# 3. ТЕХНОЛОГИИ
# ══════════════════════════════════════

add_heading_centered('3. ТЕХНОЛОГИЧЕСКИЙ СТЕК', level=1)

make_table(
    ['Компонент', 'Технология', 'Назначение'],
    [
        ['Фреймворк', 'Next.js 15 (App Router)', 'SSR/SSG, маршрутизация, API'],
        ['UI-библиотека', 'React 19 + TypeScript', 'Компоненты интерфейса'],
        ['Стили', 'Tailwind CSS 4', 'Стилизация, дизайн-система'],
        ['База данных', 'PostgreSQL 16', 'Хранение данных'],
        ['ORM', 'Prisma', 'Работа с базой данных'],
        ['Аутентификация', 'NextAuth.js v5', 'Регистрация, вход, сессии'],
        ['Оплата', 'ЮKassa API', 'Приём платежей'],
        ['Хостинг', 'VPS (Timeweb / Selectel)', 'Размещение сайта'],
        ['CDN', 'CloudFlare', 'Ускорение загрузки, защита'],
        ['Мониторинг', 'Sentry + Yandex Metrika', 'Ошибки, аналитика'],
    ],
    col_widths=[4, 5.5, 7.5]
)

# ══════════════════════════════════════
# 4. СТРУКТУРА САЙТА
# ══════════════════════════════════════

add_heading_centered('4. СТРУКТУРА САЙТА', level=1)

make_table(
    ['№', 'Раздел', 'Кол-во стр.', 'Примечание'],
    [
        ['1', 'Главная страница', '1', 'Hero, хиты, новинки, бренды'],
        ['2', 'Каталог (категории)', '15–25', 'Фильтры, сортировка, пагинация'],
        ['3', 'Карточки товаров', '478+', 'Галерея, описание, состав, цена'],
        ['4', 'Страницы брендов', '20–40', 'Логотип, описание, товары бренда'],
        ['5', 'Корзина', '1', 'Список товаров, итог'],
        ['6', 'Оформление заказа', '2', 'Данные + подтверждение'],
        ['7', 'Регистрация / Вход', '3', 'Вход, регистрация, восстановление'],
        ['8', 'Личный кабинет', '5', 'Профиль, заказы, избранное, адреса'],
        ['9', 'Информационные', '6', 'О нас, доставка, возврат, контакты, юр.'],
        ['10', 'Поиск', '1', 'Результаты поиска'],
        ['11', 'Админ-панель', '8+', 'Товары, заказы, клиенты, дашборд'],
        ['', 'ИТОГО', '540+', ''],
    ],
    col_widths=[1.5, 5.5, 3, 7]
)

# ══════════════════════════════════════
# 5. ЭТАПЫ И СРОКИ
# ══════════════════════════════════════

add_heading_centered('5. ЭТАПЫ И СРОКИ ВЫПОЛНЕНИЯ РАБОТ', level=1)

make_table(
    ['Этап', 'Содержание работ', 'Срок', 'Результат'],
    [
        ['1', 'Фундамент: инициализация проекта, база данных, '
              'дизайн-система, UI-компоненты, шапка, подвал, '
              'аутентификация',
         '7 кал. дн.', 'Работающий каркас сайта'],

        ['2', 'Каталог: страницы категорий, карточки товаров, '
              'фильтрация, поиск, страницы брендов, SEO-разметка',
         '7 кал. дн.', 'Полноценный каталог товаров'],

        ['3', 'Покупка: корзина, оформление заказа, интеграция '
              'ЮKassa, email-уведомления, информационные страницы',
         '7 кал. дн.', 'Возможность покупки и оплаты'],

        ['4', 'Личный кабинет: профиль, история заказов, '
              'избранное, адреса доставки',
         '5 кал. дн.', 'Работающий ЛК покупателя'],

        ['5', 'Главная страница, админ-панель, наполнение '
              'каталога, тестирование, оптимизация, запуск',
         '7 кал. дн.', 'Готовый к эксплуатации сайт'],
    ],
    col_widths=[1.5, 7, 3.5, 5]
)

doc.add_paragraph()
add_para('Общий срок выполнения работ: 33 (тридцать три) календарных дня '
         'с момента получения Исполнителем полного комплекта исходных материалов.', bold=True, indent_cm=1.25)

# ══════════════════════════════════════
# 6. ПОРЯДОК ОПЛАТЫ
# ══════════════════════════════════════

add_heading_centered('6. ПОРЯДОК ОПЛАТЫ', level=1)

make_table(
    ['Платёж', 'Сумма', 'Условие'],
    [
        ['Предоплата (50%)', '125 000 ₽', 'В течение 3 рабочих дней после подписания договора'],
        ['Оплата (50%)', '125 000 ₽', 'В течение 3 рабочих дней после подписания акта приёмки'],
        ['ИТОГО', '250 000 ₽', ''],
    ],
    col_widths=[5, 4, 8]
)

# ══════════════════════════════════════
# 7. ТРЕБОВАНИЯ К ПРОИЗВОДИТЕЛЬНОСТИ
# ══════════════════════════════════════

add_heading_centered('7. ТРЕБОВАНИЯ К ПРОИЗВОДИТЕЛЬНОСТИ И КАЧЕСТВУ', level=1)

make_table(
    ['Параметр', 'Требование'],
    [
        ['Время загрузки страницы', '≤ 3 секунды (при скорости соединения 10 Мбит/с)'],
        ['Google Lighthouse Score', '≥ 90 баллов (Performance, SEO, Accessibility)'],
        ['Адаптивность', 'Корректное отображение: мобильный (320px+), планшет (768px+), десктоп (1280px+)'],
        ['Кроссбраузерность', 'Chrome, Safari, Firefox, Edge — последние 2 версии'],
        ['Доступность (a11y)', 'WCAG 2.1 уровень A'],
        ['Безопасность', 'HTTPS, защита от XSS, CSRF, SQL-injection'],
        ['Uptime', '99,5% (на стороне хостинг-провайдера)'],
    ],
    col_widths=[6, 11]
)

# ══════════════════════════════════════
# 8. ИСХОДНЫЕ ДАННЫЕ
# ══════════════════════════════════════

add_heading_centered('8. ИСХОДНЫЕ ДАННЫЕ ОТ ЗАКАЗЧИКА', level=1)

add_para('Заказчик обязуется предоставить Исполнителю:', indent_cm=1.25)
add_bullet('Каталог товаров: наименования, описания, фотографии, цены, категории (478+ позиций)')
add_bullet('Логотип и элементы фирменного стиля')
add_bullet('Тексты для информационных страниц (О нас, Доставка, Возврат) или согласие на подготовку Исполнителем')
add_bullet('Доступ к хостингу и доменное имя cosmetikalux.ru')
add_bullet('Реквизиты для интеграции ЮKassa (ключи API)')
add_bullet('Доступ к Yandex Metrika (при наличии)')
add_bullet('Юридические документы: политика конфиденциальности, оферта')

doc.add_paragraph()
add_para('Срок предоставления исходных данных: не позднее 5 (пяти) рабочих дней '
         'с момента подписания договора. Задержка в предоставлении данных переносит '
         'сроки выполнения работ на соответствующий период.', indent_cm=1.25)

# ══════════════════════════════════════
# 9. КРИТЕРИИ ПРИЁМКИ
# ══════════════════════════════════════

add_heading_centered('9. КРИТЕРИИ ПРИЁМКИ', level=1)

add_heading_left('9.1. Функциональные критерии')
add_bullet('Покупатель может найти товар через каталог, фильтры или поиск')
add_bullet('Покупатель может добавить товар в корзину, оформить и оплатить заказ')
add_bullet('Покупатель может зарегистрироваться, войти и просмотреть историю заказов')
add_bullet('Администратор может добавлять, редактировать и удалять товары, категории и бренды')
add_bullet('Администратор может просматривать и обрабатывать заказы')
add_bullet('Все информационные страницы отображаются корректно')

add_heading_left('9.2. Технические критерии')
add_bullet('Lighthouse Score ≥ 90 (Performance, SEO)')
add_bullet('Время загрузки ≤ 3 секунды')
add_bullet('Корректное отображение на мобильных устройствах')
add_bullet('SEO: все страницы индексируются, Schema.org на месте, sitemap.xml генерируется')
add_bullet('HTTPS, защита от XSS/CSRF/SQL-injection')

# ══════════════════════════════════════
# 10. ДОПОЛНИТЕЛЬНЫЕ МОДУЛИ
# ══════════════════════════════════════

add_heading_centered('10. ДОПОЛНИТЕЛЬНЫЕ МОДУЛИ (НЕ ВХОДЯТ В ПАКЕТ)', level=1)

add_para('Следующие модули могут быть подключены к базовому сайту в любой момент '
         'по отдельному дополнительному соглашению:', indent_cm=1.25)

doc.add_paragraph()

make_table(
    ['Модуль', 'Стоимость', 'Срок подключения'],
    [
        ['Интеграция с 1С (синхронизация остатков, цен, заказов)', '30 000 ₽', '1 неделя'],
        ['Автоматическая доставка (СДЭК + Почта России)', '35 000 ₽', '1 неделя'],
        ['Бонусная программа (кэшбэк, уровни лояльности)', '25 000 ₽', '0,5–1 неделя'],
        ['Промокоды и акции', '20 000 ₽', '0,5 недели'],
        ['Email/SMS уведомления', '15 000 ₽', '0,5 недели'],
        ['ИИ Beauty-консультант (GigaChat)', '60 000 ₽', '2–3 недели'],
        ['Блог', '15 000 ₽', '0,5 недели'],
        ['Подписки на автозаказ', '20 000 ₽', '0,5–1 неделя'],
    ],
    col_widths=[8, 3.5, 5.5]
)

doc.add_paragraph()
add_para('Архитектура базового сайта проектируется с учётом будущего подключения '
         'всех перечисленных модулей без переработки существующего кода.', indent_cm=1.25)

# ══════════════════════════════════════
# 11. ГАРАНТИЙНЫЕ ОБЯЗАТЕЛЬСТВА
# ══════════════════════════════════════

add_heading_centered('11. ГАРАНТИЙНЫЕ ОБЯЗАТЕЛЬСТВА', level=1)

add_para('11.1. Исполнитель предоставляет гарантию на выполненные работы сроком '
         '3 (три) месяца с даты подписания акта приёмки.', indent_cm=1.25)
add_para('11.2. В течение гарантийного периода Исполнитель безвозмездно устраняет '
         'дефекты, выявленные в работе Сайта, возникшие по вине Исполнителя.', indent_cm=1.25)
add_para('11.3. Гарантия не распространяется на дефекты, возникшие в результате '
         'действий Заказчика или третьих лиц, изменений в стороннем ПО (обновления '
         'браузеров, платёжных систем и т.д.), форс-мажорных обстоятельств.', indent_cm=1.25)

# ══════════════════════════════════════
# 12. ПРАВА И ОБЯЗАННОСТИ
# ══════════════════════════════════════

add_heading_centered('12. ПЕРЕДАЧА ПРАВ', level=1)

add_para('12.1. После полной оплаты работ Исполнитель передаёт Заказчику:', indent_cm=1.25)
add_bullet('Исходный код Сайта')
add_bullet('Доступы к хостингу, базе данных, административной панели')
add_bullet('Инструкцию по управлению контентом через административную панель')

add_para('12.2. Исключительные права на разработанный программный код переходят '
         'к Заказчику с момента полной оплаты.', indent_cm=1.25)
add_para('12.3. Права на используемые библиотеки с открытым исходным кодом (open source) '
         'регулируются их собственными лицензиями.', indent_cm=1.25)

# ══════════════════════════════════════
# ПОДПИСИ
# ══════════════════════════════════════

doc.add_page_break()

add_heading_centered('ПОДПИСИ СТОРОН', level=1)

doc.add_paragraph()

table = doc.add_table(rows=8, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Убираем границы
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Убираем рамки
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), 'none')
            element.set(qn('w:sz'), '0')
            tcBorders.append(element)
        tcPr.append(tcBorders)

labels = [
    ('ЗАКАЗЧИК', 'ИСПОЛНИТЕЛЬ'),
    ('', ''),
    ('ИП Тимошенко А.В.', 'ИП Макуха Р.А.'),
    ('', ''),
    ('_______________ / А.В. Тимошенко', '_______________ / Р.А. Макуха'),
    ('', ''),
    ('«___» ____________ 2026 г.', '«___» ____________ 2026 г.'),
    ('М.П.', 'М.П.'),
]

for i, (left, right) in enumerate(labels):
    row = table.rows[i]
    for ci, text in enumerate([left, right]):
        cell = row.cells[ci]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if i == 0:
            run.bold = True

# ══════════════════════════════════════
# СОХРАНЕНИЕ
# ══════════════════════════════════════

output_path = 'ТЗ_Приложение_1_Cosmetikalux_250k.docx'
doc.save(output_path)
print(f'✓ Документ создан: {output_path}')
print(f'  Формат: DOCX (Microsoft Word)')
print(f'  Пакет: Базовый интернет-магазин — 250 000 ₽')
print(f'  Разделов: 12')
