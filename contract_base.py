#!/usr/bin/env python3
"""
Часть 1: Базовый модуль — стили, утилиты, инициализация документа.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def create_document():
    """Создать документ с настроенными стилями."""
    doc = Document()

    # Стиль Normal
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)
    style.paragraph_format.line_spacing = 1.15

    # Поля страницы
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    return doc


def add_heading_centered(doc, text, size=Pt(14), bold=True, space_before=12, space_after=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = size
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading_left(doc, text, size=Pt(12), bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = size
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_cm=1.25):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if indent_cm:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    return p


def add_item(doc, text, indent_cm=1.25):
    """Пункт с тире."""
    return add_para(doc, f'— {text}', indent_cm=indent_cm)


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def make_table(doc, headers, rows, col_widths=None, header_color='D9E2F3'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def remove_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        borders.append(el)
    tcPr.append(borders)


def make_signature_table(doc, labels):
    """Таблица подписей без рамок."""
    table = doc.add_table(rows=len(labels), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (left, right) in enumerate(labels):
        row = table.rows[i]
        for ci, text in enumerate([left, right]):
            cell = row.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if i == 0:
                run.bold = True
            remove_borders(cell)
    return table


if __name__ == '__main__':
    doc = create_document()
    add_heading_centered(doc, 'ТЕСТ')
    add_para(doc, 'Базовый модуль работает.')
    doc.save('/tmp/test_base.docx')
    print('OK: base module works')
