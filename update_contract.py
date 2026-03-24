#!/usr/bin/env python3
"""Переделка договора разработки сайта: АМИСТ (туризм) → Cosmetikalux (корейская косметика)"""

from docx import Document
from docx.shared import Pt
import copy
import re

doc = Document('Договор разработка сайта для АМИСТ_2026.docx')

# === ЗАМЕНЫ В ПАРАГРАФАХ ===

# Словарь замен текста (порядок важен — сначала длинные фразы)
replacements = [
    # Заказчик — наименование
    ('Общество с ограниченной ответственностью Тур-бизнес клуб «АМИСТ», именуемое в дальнейшем «Заказчик», в лице Генерального директора Ипатова Татьяна Николаевна, действующего на основании Устава',
     'Индивидуальный предприниматель Тимошенко Анна Вячеславовна, именуемая в дальнейшем «Заказчик», действующая на основании свидетельства о государственной регистрации'),

    # Акт приёмки
    ('ООО Тур-бизнес клуб «АМИСТ», именуемое «Заказчик», в лице Генерального директора Ипатовой Татьяны Николаевны',
     'ИП Тимошенко Анна Вячеславовна, именуемая «Заказчик»'),

    # ТЗ — заказчик
    ('ООО Тур-бизнес клуб «АМИСТ» (ИНН 6501109786, РТО 010452)',
     'ИП Тимошенко Анна Вячеславовна (ИНН 650116762461)'),

    # Предмет договора — основной
    ('оказать услуги по разработке веб-сайта для доменного имени amist.ru (далее — «Сайт»), а Заказчик обязуется принять и оплатить оказанные услуги в порядке',
     'оказать услуги по разработке интернет-магазина для доменного имени cosmetikalux.ru (далее — «Сайт»), а Заказчик обязуется принять и оплатить оказанные услуги в порядке'),

    # Результаты — базовый сайт
    ('Базовый сайт — публичная часть веб-платформы amist.ru, включающая все страницы, разделы, навигацию, формы обратной связи, поисковую оптимизацию (SEO), адаптивную вёрстку для мобильных устройств,',
     'Базовый сайт — публичная часть интернет-магазина cosmetikalux.ru, включающая каталог товаров, карточки товаров, корзину, оформление заказа, поисковую оптимизацию (SEO), адаптивную вёрстку для мобильных устройств,'),

    # 1.4 — цели сайта
    ('Сайт разрабатывается для целей продвижения услуг Заказчика в сфере туроператорской деятельности по следующим направлениям: туры по Сахалину и Курильским островам, туры в Китай (Гонконг, Пекин, Шанхай), VIP-туризм, чартерные программы, экскурсионные услуги, страхование путешественников.',
     'Сайт разрабатывается для целей продажи и продвижения товаров Заказчика в сфере корейской и японской косметики, включая: каталог продукции по категориям, карточки товаров с описаниями, корзину и оформление заказов, информационные страницы о брендах и уходе за кожей.'),

    # ТЗ — цель проекта (partial match)
    ('создание современной мультинаправленной веб-платформы для продвижения всех направлений деятельности Заказчика, обеспечивающей высокую поисковую видимость (SEO), конверсию посетителе',
     'создание современного интернет-магазина корейской и японской косметики, обеспечивающего высокую поисковую видимость (SEO), удобный каталог товаров, корзину и оформление заказов, конверсию посетителе'),

    # Название проекта в ТЗ
    ('Веб-платформа туроператора АМИСТ (amist.ru)',
     'Интернет-магазин корейской косметики Cosmetikalux (cosmetikalux.ru)'),

    # Домен
    ('amist.ru', 'cosmetikalux.ru'),

    # Подписи — Ипатова
    ('Т.Н. Ипатова', 'А.В. Тимошенко'),
    ('Ипатовой Татьяны Николаевны', 'Тимошенко Анны Вячеславовны'),
    ('Ипатова Татьяна Николаевна', 'Тимошенко Анна Вячеславовна'),

    # Email заказчика
    ('info@amist.ru', 'info@cosmetikalux.ru'),

    # Тел заказчика
    ('+7 4242 46-28-88', '+7 (962) 127-77-55'),
]


def replace_in_paragraph(paragraph, replacements):
    """Replace text in paragraph preserving formatting where possible."""
    full_text = paragraph.text
    changed = False
    for old, new in replacements:
        if old in full_text:
            full_text = full_text.replace(old, new)
            changed = True

    if changed:
        # Rebuild runs: keep first run's formatting, put all text there
        if paragraph.runs:
            fmt = paragraph.runs[0].font
            for i, run in enumerate(paragraph.runs):
                if i == 0:
                    run.text = full_text
                else:
                    run.text = ''
        else:
            paragraph.text = full_text


def replace_in_cell(cell, replacements):
    """Replace text in a table cell."""
    for paragraph in cell.paragraphs:
        replace_in_paragraph(paragraph, replacements)


# Process all paragraphs
for para in doc.paragraphs:
    replace_in_paragraph(para, replacements)

# Process all tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            replace_in_cell(cell, replacements)


# === ОБНОВЛЕНИЕ ТАБЛИЦЫ РЕКВИЗИТОВ (Table 2) ===
# Table 2 has 1 row, 2 cols: [Заказчик, Исполнитель]

req_table = doc.tables[2]
zakazchik_cell = req_table.rows[0].cells[0]

# Clear and rewrite Заказчик cell
for p in zakazchik_cell.paragraphs:
    for run in p.runs:
        run.text = ''

# Build new Заказчик text
new_zakazchik_lines = [
    'ЗАКАЗЧИК',
    '',
    'ИП Тимошенко Анна Вячеславовна',
    'ИНН: 650116762461',
    'Адрес: 693003, Сахалинская область,',
    'г. Южно-Сахалинск, пер. Земледельческий, 17',
    'Тел.: +7 (962) 127-77-55',
    'Email: info@cosmetikalux.ru',
    '',
    'Банковские реквизиты:',
    'ФИЛИАЛ "ХАБАРОВСКИЙ" АО "АЛЬФА-БАНК"',
    'БИК: 040813770',
    'К/с: 30101810800000000770',
    'Р/с: 40802810620150002751',
    '',
    'Индивидуальный предприниматель',
    '',
    '',
    '_______________ / А.В. Тимошенко',
    '',
    'М.П.',
]

# Clear all paragraphs in cell
for p in zakazchik_cell.paragraphs:
    p.clear()

# Write into first paragraph, add more as needed
paras = zakazchik_cell.paragraphs
# We need len(new_zakazchik_lines) paragraphs
while len(paras) < len(new_zakazchik_lines):
    zakazchik_cell.add_paragraph()
    paras = zakazchik_cell.paragraphs

for i, line in enumerate(new_zakazchik_lines):
    if i < len(paras):
        run = paras[i].add_run(line)
        run.font.size = Pt(10)
        if i == 0:
            run.bold = True


# === ОБНОВЛЕНИЕ СОДЕРЖАНИЯ ПОД КОСМЕТИКУ ===
# Update Table 0 (стоимость) — убрать перевод, адаптировать
table0 = doc.tables[0]
services = [
    ('Базовый интернет-магазин (каталог, карточки товаров, корзина)', '250 000', 'Этапы 1-4'),
    ('Личный кабинет (админ-панель)', '50 000', 'Этап 5'),
    ('Перевод на английский язык', '30 000', 'Этап 6 (включён)'),
    ('Перевод на китайский язык', '30 000', 'Этап 6 (включён)'),
]
for i, (name, cost, note) in enumerate(services):
    row = table0.rows[i + 1]
    row.cells[0].paragraphs[0].clear()
    row.cells[0].paragraphs[0].add_run(name).font.size = Pt(10)


# Update Table 1 (этапы) — адаптировать содержание
table1 = doc.tables[1]
stages_content = [
    ('1', 'Фундамент: дизайн-система, шапка/подвал, главная страница, SEO-основа', '4 календарных дня'),
    ('2', 'Каталог: категории товаров (478+), карточки товаров, фильтрация, поиск', '4 календарных дня'),
    ('3', 'Заказы: корзина, оформление заказа, оплата, информационные страницы', '4 календарных дня'),
    ('4', 'Интерактив: формы, аналитика, оптимизация, мобильная адаптация, запуск', '4 календарных дня'),
    ('5', 'Личный кабинет (админ-панель): управление товарами, заказами, контентом', '4 календарных дня'),
    ('6', 'Перевод на 2 иностранных языка (английский и китайский)', '7 календарных дня'),
]
for i, (num, content, term) in enumerate(stages_content):
    row = table1.rows[i + 1]
    row.cells[1].paragraphs[0].clear()
    row.cells[1].paragraphs[0].add_run(content).font.size = Pt(10)


# Update Table 3 (разделы сайта) — адаптировать структуру
table3 = doc.tables[3]
sections = [
    ('Главная страница', '1', ''),
    ('Каталог товаров (категории)', '15', ''),
    ('Карточки товаров', '478+', ''),
    ('Корзина и оформление заказа', '3', ''),
    ('О магазине / О нас', '2', ''),
    ('Бренды', '10-15', ''),
    ('Блог / Уход за кожей', '5-10', ''),
    ('Доставка и оплата', '2', ''),
    ('Контакты', '1', ''),
    ('Юридические страницы', '4', 'Политика, оферта, оплата, реквизиты'),
    ('Акции / Новинки', '2', ''),
    ('FAQ', '1', ''),
    ('Прочие', '2-4', ''),
]

# We have 15 rows (header + 14 data). Need to adjust.
# Clear existing rows and fill with new data
for i in range(1, len(table3.rows)):  # skip header
    row = table3.rows[i]
    if i - 1 < len(sections):
        name, count, note = sections[i - 1]
        for ci, val in enumerate([name, count, note]):
            row.cells[ci].paragraphs[0].clear()
            row.cells[ci].paragraphs[0].add_run(val).font.size = Pt(10)
    elif i == len(table3.rows) - 1:  # last row = ИТОГО
        row.cells[0].paragraphs[0].clear()
        row.cells[0].paragraphs[0].add_run('ИТОГО').font.size = Pt(10)
        run = row.cells[0].paragraphs[0].runs[0]
        run.bold = True
        row.cells[1].paragraphs[0].clear()
        row.cells[1].paragraphs[0].add_run('~530').font.size = Pt(10)
        row.cells[2].paragraphs[0].clear()


# Update Table 6 (этапы в ТЗ)
table6 = doc.tables[6]
tz_stages = [
    ('1', 'Фундамент: дизайн-система, шапка, подвал, главная', '28 кал. дн.', 'Работающая главная страница'),
    ('2', 'Каталог: категории, карточки товаров (478+), фильтры, поиск', '35 кал. дн.', 'Полный каталог товаров'),
    ('3', 'Заказы: корзина, оформление, оплата, информационные страницы', '28 кал. дн.', 'Функционал заказов'),
    ('4', 'Интерактив, формы, аналитика, оптимизация, запуск', '28 кал. дн.', 'Готовый к запуску магазин'),
    ('5', 'Личный кабинет (админ-панель)', '28 кал. дн.', 'Работающая админ-панель'),
    ('6', 'Перевод на английский и китайский языки', '28 кал. дн.', 'Мультиязычный сайт'),
]
for i, (num, content, term, result) in enumerate(tz_stages):
    row = table6.rows[i + 1]
    row.cells[1].paragraphs[0].clear()
    row.cells[1].paragraphs[0].add_run(content).font.size = Pt(10)
    row.cells[3].paragraphs[0].clear()
    row.cells[3].paragraphs[0].add_run(result).font.size = Pt(10)


# Update Table 7 (календарный план) — содержание этапов
table7 = doc.tables[7]
cal_stages = [
    ('1', 'Фундамент и главная страница'),
    ('2', 'Каталог товаров (категории, карточки, фильтры, поиск)'),
    ('3', 'Корзина, заказы, информационные страницы'),
    ('4', 'Интерактив, оптимизация, запуск'),
    ('5', 'Личный кабинет (админ-панель)'),
]
for i, (num, content) in enumerate(cal_stages):
    row = table7.rows[i + 1]
    row.cells[1].paragraphs[0].clear()
    row.cells[1].paragraphs[0].add_run(content).font.size = Pt(10)


# === ЗАМЕНЫ В РАЗДЕЛЕ «Исходные данные» ===
for para in doc.paragraphs:
    text = para.text
    if 'данные 25 туров, 46 экскурсий' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '6.1. Заказчик предоставляет: данные 478 товаров (фото, описания, цены, категории), логотип и фирменный стиль, юридические документы, доступы к хостингу и Yandex Metrika.'
    elif 'контент разделов Китай' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '6.2. Исполнитель создаёт: дизайн-систему магазина, шаблоны страниц, SEO-тексты и мета-теги для всех страниц, информационный контент (уход за кожей, бренды).'
    # ТЗ функциональные требования — публичная часть
    elif 'Schema.org: TravelAgency' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '   — Schema.org: Product, Offer, Organization, BreadcrumbList, WebSite;'
    elif 'фотогалереи с лайтбоксом и свайп-навигацией' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '   — каталог товаров с фильтрацией по категориям, брендам и цене;'
    elif 'формы обратной связи и бронирования с валидацией' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '   — корзина и оформление заказа с валидацией;'
    elif 'виджет ИИ-ассистента (UI-заглушка для GigaChat)' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '   — карточки товаров с фото, описанием, ценой и кнопкой «В корзину»;'
    # Админ-панель
    elif 'CRUD для всех типов контента (туры, экскурсии, блог, страницы)' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '   — CRUD для всех типов контента (товары, категории, бренды, страницы);'
    elif 'управление заявками: список, статусы, фильтры, экспорт' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '   — управление заказами: список, статусы, фильтры, экспорт;'
    # Личный кабинет
    elif 'дашборд с KPI (посещения, заявки, конверсия)' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '   — дашборд с KPI (посещения, заказы, конверсия, средний чек);'
    # Перевод
    elif 'создание страниц привлечения иностранных турагентов' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '   — создание мультиязычных карточек товаров;'
    # Общее описание результатов
    elif 'Личный кабинет (админ-панель) — административный интерфейс для управления контентом сайта, заявками клиентов' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '   б) Личный кабинет (админ-панель) — административный интерфейс для управления каталогом товаров, заказами клиентов, медиа-библиотекой, аналитикой и настройками, включающий систему ролей (Администратор, Редактор, Менеджер);'


# Schema for ТЗ table 4 (tech stack) — update AI assistant row
table4 = doc.tables[4]
for row in table4.rows:
    if 'ИИ-ассистент' in row.cells[0].text:
        row.cells[0].paragraphs[0].clear()
        row.cells[0].paragraphs[0].add_run('Корзина/заказы').font.size = Pt(10)
        row.cells[1].paragraphs[0].clear()
        row.cells[1].paragraphs[0].add_run('Сессии, localStorage, серверная валидация').font.size = Pt(10)


# Save
output_path = 'Договор разработка сайта для Cosmetikalux_2026.docx'
doc.save(output_path)
print(f'Сохранено: {output_path}')
